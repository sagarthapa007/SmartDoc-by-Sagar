import csv
import io
from typing import Any, Dict, List, Optional

from app.services.registry import REGISTRY

try:
    import pandas as pd
except Exception:  # If pandas not available at runtime, provide minimal CSV-only path
    pd = None
try:
    import docx  # python-docx
except Exception:
    docx = None


def parse_csv_bytes(raw: bytes) -> Dict[str, Any]:
    text = raw.decode("utf-8", errors="ignore")
    reader = csv.DictReader(io.StringIO(text))
    rows = list(reader)
    headers = reader.fieldnames or (rows[0].keys() if rows else [])
    return {"headers": list(headers or []), "rows": rows}


def parse_xlsx_bytes(raw: bytes) -> Dict[str, Any]:
    if pd is None:
        raise RuntimeError("pandas is required for Excel parsing")
    bio = io.BytesIO(raw)
    df = pd.read_excel(bio)
    headers = list(df.columns.astype(str))
    rows = df.astype(object).where(pd.notnull(df), None).to_dict(orient="records")
    return {"headers": headers, "rows": rows}


def parse_docx_bytes(raw: bytes) -> Dict[str, Any]:
    if docx is None:
        # Fallback: treat as plain text
        text = raw.decode("utf-8", errors="ignore")
        return {"headers": [], "rows": [], "text_blocks": [text]}
    bio = io.BytesIO(raw)
    document = docx.Document(bio)
    text_blocks: List[str] = []
    tables_data: List[Dict[str, Any]] = []

    # Collect paragraphs
    for p in document.paragraphs:
        if p.text and p.text.strip():
            text_blocks.append(p.text.strip())

    # Collect first table as dataset if present
    headers, rows = [], []
    for tbl in document.tables:
        # Try first row as headers
        hdr = [cell.text.strip() for cell in tbl.rows[0].cells]
        headers = hdr if any(hdr) else []
        for r in tbl.rows[1:]:
            row = [cell.text.strip() for cell in r.cells]
            if headers and len(row) == len(headers):
                rows.append(dict(zip(headers, row)))
        if rows:
            break

    result = {"headers": headers, "rows": rows}
    if text_blocks:
        result["text_blocks"] = text_blocks
    return result


def parse_upload(filename: str, raw: bytes) -> Dict[str, Any]:
    lower = filename.lower()
    if lower.endswith(".csv"):
        return parse_csv_bytes(raw)
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        return parse_xlsx_bytes(raw)
    if lower.endswith(".docx") or lower.endswith(".doc"):
        return parse_docx_bytes(raw)
    if lower.endswith(".txt"):
        text = raw.decode("utf-8", errors="ignore")
        return {"headers": [], "rows": [], "text_blocks": [text]}
    # Default: try CSV

    return parse_csv_bytes(raw)
